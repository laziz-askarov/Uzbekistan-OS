#!/usr/bin/env python3
"""Build the reviewed Uzbekistan OS visa knowledge index from retained DOCX sources."""

from __future__ import annotations

import argparse
import hashlib
import json
from datetime import date
from pathlib import Path

from docx import Document


DOCUMENTS = {
    "Business visa.docx": ("business-visa", "Business visa", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "e-visa UZ.docx": ("electronic-visa", "Electronic visa", "https://www.e-visa.gov.uz/"),
    "Family reunification.docx": ("family-visit", "Family reunification and private visits", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "LIST_OF_CATEGORIES_OF_ENTRY,_EXIT_AND_TRANSIT_VISAS_NON_ELECTRONIC.docx": ("visa-categories", "Non-electronic visa category catalogue", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Overstay penalties.docx": ("overstay-and-exit", "Overstay penalties and exit", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Passport validity requirements.docx": ("passport-validity", "Passport validity requirements", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Permanent Residence (PR) in Uzbekistan.docx": ("permanent-residence", "Permanent residence in Uzbekistan", "https://my.gov.uz/en/for-foreigners"),
    "Registration of foreigners.docx": ("arrival-registration", "Registration of foreigners", "https://my.gov.uz/en/for-foreigners"),
    "Student visas.docx": ("student-visa", "Student visas", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Temporary residence permits.docx": ("temporary-residence", "Temporary residence permits", "https://my.gov.uz/en/for-foreigners"),
    "Visa categories.docx": ("visa-categories", "Visa categories", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Visa to the Republic of Uzbekistan MFA.docx": ("mfa-visa-guidance", "Visa to the Republic of Uzbekistan — MFA", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
    "Visa-free entry to Uzbekistan.docx": ("visa-free-entry", "Visa-free entry to Uzbekistan", "https://gov.uz/en/mfa/activity_page/o-zbekiston-respublikasi-vizasi"),
}


def extract_blocks(path: Path) -> list[str]:
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return blocks


def chunk_blocks(blocks: list[str], target_chars: int = 1_800) -> list[str]:
    chunks: list[str] = []
    current: list[str] = []
    current_size = 0
    for block in blocks:
        if current and current_size + len(block) + 2 > target_chars:
            chunks.append("\n\n".join(current))
            current = current[-1:]
            current_size = sum(len(item) + 2 for item in current)
        current.append(block)
        current_size += len(block) + 2
    if current:
        chunks.append("\n\n".join(current))
    return chunks


def build_index(source_dir: Path) -> dict[str, object]:
    imported_at = date.today().isoformat()
    sources: list[dict[str, object]] = []
    chunks: list[dict[str, str]] = []
    for filename, (topic, title, url) in DOCUMENTS.items():
        path = source_dir / filename
        if not path.exists():
            raise FileNotFoundError(f"Missing required source: {path}")
        data = path.read_bytes()
        digest = hashlib.sha256(data).hexdigest()
        blocks = extract_blocks(path)
        document_chunks = chunk_blocks(blocks)
        sources.append(
            {
                "filename": filename,
                "title": title,
                "topic": topic,
                "officialUrl": url,
                "sha256": digest,
                "blockCount": len(blocks),
                "chunkCount": len(document_chunks),
            }
        )
        for index, content in enumerate(document_chunks, start=1):
            chunks.append(
                {
                    "id": f"{topic}-{hashlib.sha256(filename.encode()).hexdigest()[:8]}-{index}",
                    "title": title,
                    "topic": topic,
                    "url": url,
                    "reviewedAt": imported_at,
                    "sourceFile": filename,
                    "sourceSha256": digest,
                    "content": content,
                }
            )
    return {
        "schemaVersion": 1,
        "importedAt": imported_at,
        "sourceCount": len(sources),
        "chunkCount": len(chunks),
        "sources": sources,
        "chunks": chunks,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()
    index = build_index(args.source_dir)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(index, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"Indexed {index['sourceCount']} documents into {index['chunkCount']} chunks")


if __name__ == "__main__":
    main()
